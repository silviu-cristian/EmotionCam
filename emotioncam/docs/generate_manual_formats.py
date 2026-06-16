"""Generate the offline EmotionCam DOCX manual from the maintained Markdown source."""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE = DOCS / "user_manual.md"
OUTPUT = DOCS / "EmotionCam_User_Manual.docx"
ICON = ROOT / "app" / "assets" / "icon.png"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(ICON), width=Inches(1.45))
    paragraph.paragraph_format.space_before = Pt(55)
    paragraph.paragraph_format.space_after = Pt(22)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EmotionCam User Manual")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string("203748")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Local-first visible-expression estimates with optional Local Ollama or OpenAI AI analysis")
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("52687C")
    date = document.add_paragraph("Version 1.2.0 AI-enabled | Updated June 16, 2026")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [[item.strip() for item in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", item) for item in rows[1]):
        rows.pop(1)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    table.style = "Table Grid"
    width = Inches(6.5 / len(rows[0]))
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = value.replace("`", "")
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def add_body(document: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["Normal"]
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        image = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image:
            path = (DOCS / image.group(2)).resolve()
            if path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(path), width=Inches(5.9))
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        elif re.match(r"^\d+\.\s", stripped):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(re.sub(r"^\d+\.\s+", "", stripped).replace("**", "").replace("`", ""))
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped[2:].replace("**", "").replace("`", ""))
        elif stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(stripped[2:].replace("**", "").replace("`", ""))
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("7A5A00")
        elif stripped:
            document.add_paragraph(stripped.replace("**", "").replace("`", ""))
        index += 1


document = Document()
style_document(document)
add_cover(document)
add_body(document)
footer = document.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("EmotionCam User Manual | Local visible-expression estimates")
document.save(OUTPUT)
print(OUTPUT)
